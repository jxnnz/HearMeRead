"""
Utility functions for parsing uploaded documents into passage content and questions.

Supported formats:
  - .docx  — parsed with python-docx (free, no external APIs)
  - .txt   — parsed as plain text

Combined format (single file for both passage + questions):
─────────────────────────────────────────────────────────
  [PASSAGE]
  Noong unang panahon, may isang pagong...

  [QUESTIONS]
  Sino ang pangunahing tauhan?
  Ano ang aral ng kwento?
─────────────────────────────────────────────────────────

Section markers are case-insensitive. Numbering prefixes on questions
(1. / 2) / Q1: / (1) / Q1.) are stripped automatically.
"""

import re
from dataclasses import dataclass
from io import BytesIO
from typing import List, Optional, Tuple

from docx import Document
from fastapi import HTTPException, status


# Constants
_PASSAGE_MARKER   = re.compile(r"^\[PASSAGE\]\s*$",   re.IGNORECASE)
_QUESTIONS_MARKER = re.compile(r"^\[QUESTIONS\]\s*$", re.IGNORECASE)
_NUMBER_PREFIX    = re.compile(r"^\s*(?:Q\s*)?\d+[\.\)\-\:]\s*|^\s*\(\d+\)\s*", re.IGNORECASE)

_MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB


# Result dataclass
@dataclass
class ParsedDocument:
    passage_content: str
    questions: List[str]


# Internal helpers
def _strip_numbering(text: str) -> str:
    return _NUMBER_PREFIX.sub("", text).strip()


def _split_sections(lines: List[str]) -> Tuple[List[str], List[str]]:
    """
    Given a flat list of text lines, split into passage lines and question lines
    using [PASSAGE] and [QUESTIONS] markers.
    Raises HTTP 422 if either section marker is missing.
    """
    passage_lines:  List[str] = []
    question_lines: List[str] = []
    current_section: Optional[str] = None

    for line in lines:
        stripped = line.strip()
        if _PASSAGE_MARKER.match(stripped):
            current_section = "passage"
            continue
        if _QUESTIONS_MARKER.match(stripped):
            current_section = "questions"
            continue

        if current_section == "passage" and stripped:
            passage_lines.append(stripped)
        elif current_section == "questions" and stripped:
            question_lines.append(stripped)

    if not passage_lines and not question_lines:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "Could not find [PASSAGE] or [QUESTIONS] markers in the document. "
                "Please follow the required template format."
            ),
        )
    if not passage_lines:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="[PASSAGE] section is empty or marker is missing.",
        )
    if not question_lines:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="[QUESTIONS] section is empty or marker is missing.",
        )

    return passage_lines, question_lines


# File readers
def _read_docx_lines(file_bytes: bytes) -> List[str]:
    try:
        doc = Document(BytesIO(file_bytes))
        return [p.text for p in doc.paragraphs]
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not read the file. Make sure it is a valid .docx document.",
        )


def _read_txt_lines(file_bytes: bytes) -> List[str]:
    try:
        return file_bytes.decode("utf-8").splitlines()
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not decode the .txt file. Make sure it is saved as UTF-8.",
        )


def _get_lines(file_bytes: bytes, filename: str) -> List[str]:
    name = filename.lower()
    if name.endswith(".docx"):
        return _read_docx_lines(file_bytes)
    if name.endswith(".txt"):
        return _read_txt_lines(file_bytes)
    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail="Only .docx and .txt files are accepted.",
    )


# Public API
def validate_upload(file_bytes: bytes, filename: str) -> None:
    """Enforce size and type limits before parsing."""
    if len(file_bytes) > _MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File exceeds the 5 MB size limit.",
        )
    name = filename.lower()
    if not (name.endswith(".docx") or name.endswith(".txt")):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Only .docx and .txt files are accepted.",
        )


def parse_combined(file_bytes: bytes, filename: str) -> ParsedDocument:
    """
    Parse a combined document with [PASSAGE] and [QUESTIONS] sections.
    Returns passage content string + list of question strings.
    """
    lines = _get_lines(file_bytes, filename)
    passage_lines, question_lines = _split_sections(lines)

    passage_content = "\n".join(passage_lines)
    questions = [_strip_numbering(q) for q in question_lines if _strip_numbering(q)]

    return ParsedDocument(
        passage_content=passage_content,
        questions=questions,
    )


def parse_passage_only(file_bytes: bytes, filename: str) -> str:
    """
    Parse a document containing only passage text (no section markers needed).
    Every non-empty paragraph/line is joined as the passage content.
    """
    lines = _get_lines(file_bytes, filename)
    content_lines = [l.strip() for l in lines if l.strip()]

    if not content_lines:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="The uploaded document appears to be empty.",
        )

    return "\n".join(content_lines)


def parse_questions_only(file_bytes: bytes, filename: str) -> List[str]:
    """
    Parse a document containing only questions (no section markers needed).
    Each non-empty paragraph/line becomes one question.
    Numbering prefixes are stripped automatically.
    """
    lines = _get_lines(file_bytes, filename)
    questions = [_strip_numbering(l.strip()) for l in lines if l.strip()]
    questions = [q for q in questions if q]

    if not questions:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No questions found in the uploaded document.",
        )

    return questions